import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx package is required. Install it with: pip install python-docx")
    raise SystemExit(1)

def create_user_document():
    doc = Document()

    # Title
    title = doc.add_heading('Campus Placement System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('User Documentation', 1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Objective
    doc.add_heading('1. Objective of the System', level=1)
    p = doc.add_paragraph(
        "The Campus Placement System is a comprehensive, multi-tenant SaaS platform designed to bridge the gap between educational institutions, students, and employers. "
        "The primary objective is to automate and streamline the end-to-end recruitment process within a college or university. "
        "By providing a centralized hub for communication, application tracking, and placement drive management, the system ensures transparency, efficiency, and better outcomes for all stakeholders."
    )

    # 2. Benefits
    doc.add_heading('2. Key Benefits', level=1)
    
    doc.add_heading('For Students:', level=2)
    doc.add_paragraph('Accessible job/internship listings and simplified application process.', style='List Bullet')
    doc.add_paragraph('Real-time tracking of application status and personal interview schedules.', style='List Bullet')
    doc.add_paragraph('Digital professional profile and resume management.', style='List Bullet')

    doc.add_heading('For Employers:', level=2)
    doc.add_paragraph('Streamlined candidate sourcing and screening tools.', style='List Bullet')
    doc.add_paragraph('Efficient management of placement drives and interview logistics.', style='List Bullet')
    doc.add_paragraph('Direct collaboration with college placement cells.', style='List Bullet')

    doc.add_heading('For Colleges (TPO):', level=2)
    doc.add_paragraph('Centralized student data and employer relationship management.', style='List Bullet')
    doc.add_paragraph('Automated record-keeping and data-driven placement reports.', style='List Bullet')
    doc.add_paragraph('Simplified scheduling of drives and infrastructure allocation.', style='List Bullet')

    # 3. Features & Screens
    doc.add_heading('3. Features and Screens', level=1)

    # Student Module
    doc.add_heading('A. Student Module', level=2)
    features_student = {
        "Dashboard": "Central overview of active applications, upcoming drive alerts, and latest news.",
        "Applications": "Detailed view of all jobs applied to, including history and current status.",
        "Jobs & Internships": "Search and apply for various career opportunities.",
        "Drives": "Calendar view of upcoming on-campus and off-campus recruitment events.",
        "Interviews": "Personal schedule for upcoming selection rounds.",
        "Offers": "Management of job and internship offer letters.",
        "Profile": "Digital resume and portfolio including skills, academics, and projects.",
        "Discussions": "Forums for peer-to-peer interaction and TPO announcements.",
        "Documents": "Secure storage for academic and professional certificates."
    }
    for feature, desc in features_student.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{feature}: ").bold = True
        p.add_run(desc)

    # Employer Module
    doc.add_heading('B. Employer Module', level=2)
    features_employer = {
        "Dashboard": "Analytics on recruitment progress and active vacancy status.",
        "Job/Internship Management": "Interface for posting and managing job requirements.",
        "Application Review": "Tools for screening candidates and managing the shortlisting process.",
        "Interview Scheduling": "Collaboration tool for scheduling student interviews.",
        "Placement Drives": "Dashboard for managing participation in specific college drives.",
        "Offer Release": "System to generate and track offers sent to selected students.",
        "Projects & Sponsorships": "Manage collaborative projects and campus branding initiatives.",
        "Select Campus": "Tool to target and manage relationships with multiple colleges.",
        "Company Profile": "Identity management for presenting the organization to students."
    }
    for feature, desc in features_employer.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{feature}: ").bold = True
        p.add_run(desc)

    # College Admin Module
    doc.add_heading('C. College Admin (TPO) Module', level=2)
    features_college = {
        "Dashboard": "High-level overview of placement stats and student eligibility status.",
        "Student Management": "Comprehensive database for monitoring and managing student records.",
        "Employer Relations": "CRM system to manage company partnerships and history.",
        "Drive Coordination": "End-to-end event planning for placement recruitment drives.",
        "Infrastructure": "Management of physical resources like halls and labs for drives.",
        "Reports & Analytics": "Automated generation of placement success and data reports.",
        "Rules & Guidelines": "Policy setting for placement eligibility and fair practices.",
        "Calendar & Events": "Master schedule of all activities within the placement cell.",
        "Feedback & Alerts": "System for gathering stakeholder feedback and broadcasting urgent notifications."
    }
    for feature, desc in features_college.items():
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{feature}: ").bold = True
        p.add_run(desc)

    # Save the document
    default_out = 'docs/product/campus-placement-system-user-document.docx'
    file_path = sys.argv[1] if len(sys.argv) > 1 else default_out
    try:
        doc.save(file_path)
        print(f"Document saved successfully as {file_path}")
    except OSError as e:
        print(f"Error saving document to {file_path}: {e}")
        raise SystemExit(1)

if __name__ == "__main__":
    create_user_document()
