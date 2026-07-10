"""Convert PlacementHub functionality markdown to Word (.docx) with embedded diagram PNGs."""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from render_mermaid_png import DIAGRAMS_DIR, diagram_slug, ensure_diagram_pngs, extract_mermaid_blocks


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def strip_md_links(text: str) -> str:
    """Convert [label](url) to label (url) for plain Word text."""
    text = re.sub(r"\[`([^`]+)`\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def add_rich_paragraph(doc, text: str, style=None):
    """Paragraph with **bold** inline."""
    text = strip_md_links(text)
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            subparts = re.split(r"(`[^`]+`)", part)
            for sp in subparts:
                if sp.startswith("`") and sp.endswith("`"):
                    run = p.add_run(sp[1:-1])
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                else:
                    p.add_run(sp)
    return p


def add_diagram_image(doc, png_path: Path, *, width_in: float = 6.5) -> None:
    if not png_path.is_file():
        doc.add_paragraph(f"[Diagram missing: {png_path.name}]", style="Intense Quote")
        return
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    run.add_picture(str(png_path), width=Inches(width_in))
    doc.add_paragraph()


def parse_table_row(line: str):
    line = line.strip()
    if not line.startswith("|"):
        return None
    return [c.strip() for c in line.strip("|").split("|")]


def is_separator_row(cells):
    return all(re.match(r"^:?-+:?$", c.replace(" ", "")) for c in cells if c)


def build_diagram_map(md_text: str) -> dict[str, Path]:
    mapping: dict[str, Path] = {}
    for i, block in enumerate(extract_mermaid_blocks(md_text), start=1):
        slug = diagram_slug(block, i)
        mapping[block] = DIAGRAMS_DIR / f"{slug}.png"
    return mapping


def convert_md_to_docx(md_path: Path, docx_path: Path, *, force_diagrams: bool = False) -> None:
    md_text = md_path.read_text(encoding="utf-8")
    lines = md_text.splitlines()

    ensure_diagram_pngs(md_path, force=force_diagrams)
    diagram_by_source = build_diagram_map(md_text)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title_set = False
    i = 0
    in_code = False
    code_lines: list[str] = []
    in_mermaid = False
    mermaid_lines: list[str] = []
    skip_meta_line = False

    while i < len(lines):
        line = lines[i]

        # Skip Word-export meta line in header
        if line.strip().startswith("**Word export:**"):
            i += 1
            continue

        if line.strip().startswith("```mermaid"):
            in_mermaid = True
            mermaid_lines = []
            i += 1
            continue

        if in_mermaid:
            if line.strip() == "```":
                source = "\n".join(mermaid_lines).strip()
                png = diagram_by_source.get(source)
                if png:
                    # Wider diagram for horizontal flow; assessment flow is taller
                    width = 6.8 if "flowchart LR" in source else 5.5
                    add_diagram_image(doc, png, width_in=width)
                else:
                    doc.add_paragraph("[Diagram could not be rendered.]", style="Intense Quote")
                in_mermaid = False
                mermaid_lines = []
            else:
                mermaid_lines.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph("\n".join(code_lines))
                p.style = "Intense Quote"
                for run in p.runs:
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if line.strip().startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = parse_table_row(lines[i])
                if cells and not is_separator_row(cells):
                    table_rows.append(cells)
                i += 1
            if table_rows:
                ncols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=ncols)
                table.style = "Table Grid"
                for ri, row in enumerate(table_rows):
                    for ci in range(ncols):
                        cell_text = row[ci] if ci < len(row) else ""
                        cell_text = strip_md_links(re.sub(r"\*\*([^*]+)\*\*", r"\1", cell_text))
                        cell = table.rows[ri].cells[ci]
                        cell.text = cell_text
                        if ri == 0:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.bold = True
                            set_cell_shading(cell, "E8EEF7")
                doc.add_paragraph()
            continue

        if line.startswith("# "):
            if not title_set:
                h = doc.add_heading(strip_md_links(line[2:].strip()), level=0)
                h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                title_set = True
            else:
                doc.add_heading(strip_md_links(line[2:].strip()), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(strip_md_links(line[3:].strip()), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(strip_md_links(line[4:].strip()), level=3)
            i += 1
            continue

        if re.match(r"^-\s+", line):
            add_rich_paragraph(doc, re.sub(r"^-\s+", "", line), style="List Bullet")
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)", line)
        if m:
            add_rich_paragraph(doc, m.group(2), style="List Number")
            i += 1
            continue

        if line.strip().startswith("*") and line.strip().endswith("*") and not line.strip().startswith("**"):
            p = doc.add_paragraph(strip_md_links(line.strip()[1:-1]))
            p.runs[0].italic = True
            i += 1
            continue

        if line.strip():
            add_rich_paragraph(doc, line.strip())
        i += 1

    doc.add_paragraph()
    fp = doc.add_paragraph("PlacementHub — Functionality, Features & Flows · June 2026 (sandbox / pre-launch)")
    fp.runs[0].font.size = Pt(9)
    fp.runs[0].italic = True

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "docs" / "product" / "placementhub-functionality.md"
    out = root / "docs" / "product" / "placementhub-functionality.docx"
    force = "--force-diagrams" in sys.argv
    if len(sys.argv) > 1 and not sys.argv[1].startswith("--"):
        out = Path(sys.argv[1])
    convert_md_to_docx(md, out, force_diagrams=force)
